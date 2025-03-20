import os
import PyPDF2
import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io

def read_pdf(file_path):
    """
    Extract text from a PDF file.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text from the PDF
    """
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def read_docx(file_path):
    """
    Extract text from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the DOCX
    """
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

# Note: read_template function is now in template_manager.py

def fill_template(template_text, data):
    """
    Fill a template with data.
    
    Args:
        template_text (str): The text content of the template
        data (dict): A dictionary with field names and values
        
    Returns:
        str: Filled template
    """
    filled_template = template_text
    for field, value in data.items():
        placeholder = f"[{field}]"
        filled_template = filled_template.replace(placeholder, value)
    return filled_template

def generate_pdf(text, output_path):
    """
    Generate a PDF from text.
    
    Args:
        text (str): Text to include in the PDF
        output_path (str): Path to save the PDF
        
    Returns:
        bool: Success status
    """
    try:
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=letter)
        
        # Set up text formatting
        c.setFont("Helvetica", 12)
        
        # Calculate page width and height
        width, height = letter
        
        # Start at the top of the page
        y_position = height - 50
        
        # Split text into lines and write to PDF
        for line in text.split('\n'):
            if y_position < 50:  # Create a new page if we're at the bottom
                c.showPage()
                c.setFont("Helvetica", 12)
                y_position = height - 50
            
            c.drawString(50, y_position, line)
            y_position -= 15
        
        c.save()
        
        # Get the value of the BytesIO buffer and write it to a file
        packet.seek(0)
        with open(output_path, 'wb') as f:
            f.write(packet.getvalue())
            
        return True
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return False

def generate_docx(text, output_path):
    """
    Generate a DOCX from text.
    
    Args:
        text (str): Text to include in the DOCX
        output_path (str): Path to save the DOCX
        
    Returns:
        bool: Success status
    """
    try:
        doc = docx.Document()
        
        # Split text into paragraphs and add to document
        for paragraph in text.split('\n'):
            doc.add_paragraph(paragraph)
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        return False 