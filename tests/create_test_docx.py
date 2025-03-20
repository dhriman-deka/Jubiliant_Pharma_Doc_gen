"""
Create a test DOCX document for testing purposes.
"""
import os
import docx
from pathlib import Path

# Get the directory of this script
SCRIPT_DIR = Path(__file__).parent
OUTPUT_DIR = SCRIPT_DIR / "documents"

# Create output directory if it doesn't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Create a new document
doc = docx.Document()

# Add a heading
doc.add_heading('TEST PROPOSAL DOCUMENT', 0)

# Add company information
doc.add_paragraph('Tech Innovations LLC')
doc.add_paragraph('789 Innovation Drive')
doc.add_paragraph('San Francisco, CA 94105')
doc.add_paragraph('Contact: Mark Johnson')
doc.add_paragraph('Email: mjohnson@techinnovations.example')
doc.add_paragraph('Phone: (555) 987-6543')

# Add a section break
doc.add_paragraph('').add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

# Add proposal details
doc.add_heading('Project Proposal: Mobile App Development', level=1)
doc.add_paragraph('Client: Global Retail Solutions')
doc.add_paragraph('Date: August 15, 2023')

# Add project description
doc.add_heading('Project Description', level=2)
description = doc.add_paragraph()
description.add_run('Tech Innovations LLC proposes to develop a mobile application for Global Retail Solutions that will enhance customer engagement and streamline the shopping experience. The application will include the following features:')

# Add bullet points
features = [
    'User account management and authentication',
    'Product browsing and search functionality',
    'Shopping cart and checkout process',
    'Order tracking and history',
    'Customer support chat integration',
    'Push notifications for promotions and updates'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# Add timeline section
doc.add_heading('Project Timeline', level=2)
doc.add_paragraph('Start Date: September 1, 2023')
doc.add_paragraph('End Date: February 28, 2024')

# Add pricing section
doc.add_heading('Pricing', level=2)
doc.add_paragraph('Total Project Cost: $120,000')
doc.add_paragraph('Payment Schedule:')

# Add payment schedule
payments = [
    '25% upon contract signing ($30,000)',
    '25% after completion of UI/UX design phase ($30,000)',
    '25% after development of core functionality ($30,000)',
    '25% upon project completion and acceptance ($30,000)'
]

for payment in payments:
    doc.add_paragraph(payment, style='List Bullet')

# Add terms and conditions
doc.add_heading('Terms and Conditions', level=2)
doc.add_paragraph('All work will be performed according to the specifications provided in the attached Statement of Work.')
doc.add_paragraph('Changes to the project scope may result in additional costs and timeline adjustments.')
doc.add_paragraph('Tech Innovations LLC retains ownership of all intellectual property until final payment is received.')

# Add signature section
doc.add_heading('Signatures', level=2)
doc.add_paragraph('_______________________________     _________________')
doc.add_paragraph('Mark Johnson, CEO                    Date')
doc.add_paragraph('Tech Innovations LLC')
doc.add_paragraph('\n')
doc.add_paragraph('_______________________________     _________________')
doc.add_paragraph('Client Representative                Date')
doc.add_paragraph('Global Retail Solutions')

# Save the document
doc_path = OUTPUT_DIR / 'test_proposal.docx'
doc.save(str(doc_path))

print(f"Test DOCX document created at: {doc_path}")

# Create a PDF placeholder
pdf_path = OUTPUT_DIR / 'test_document.pdf'
with open(pdf_path, 'w') as f:
    f.write("This is a placeholder for a PDF test document. In a real implementation, you would use a PDF generation library.")

print(f"PDF placeholder created at: {pdf_path}") 